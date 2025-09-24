import os
import json
from dotenv import load_dotenv
from typing import TypedDict, List, Dict, Any
from langgraph.graph import StateGraph, END
from langchain_openai import AzureChatOpenAI
from docx import Document as DocxDocument
import fitz  # PyMuPDF for PDF reading
import logging
import time
from md2docx_python.src.md2docx_python import markdown_to_word
import streamlit as st
import base64

load_dotenv()

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s - %(message)s",
    handlers=[
        logging.FileHandler("testgenerationagent.log"),  # Logs to file
        logging.StreamHandler()               # Logs to console
    ]
)

logger = logging.getLogger("TestScriptGenerationAgent")

class StreamlitLogger(logging.Handler):
    def __init__(self, placeholder):
        super().__init__()
        self.placeholder = placeholder
        self.logs = []

    def emit(self, record):
        log_entry = self.format(record)
        self.logs.append(log_entry)
        # Display logs in real-time
        self.placeholder.text("\n".join(self.logs[-15:]))  # keep last 15 lines visible

def telemetry(node_name):
    def decorator(func):
        def wrapper(state, *args, **kwargs):
            start = time.time()
            logger.info(f"Starting node: {node_name}")
            try:
                result = func(state, *args, **kwargs)
                duration = time.time() - start
                logger.info(f"Completed node: {node_name} in {duration:.2f}s")
                return result
            except Exception as e:
                duration = time.time() - start
                logger.error(f"Failed node: {node_name} in {duration:.2f}s with error: {e}")
                raise
        return wrapper
    return decorator

# -------------------------------
# 1. Agent State
# -------------------------------
class AgentState(TypedDict, total=False):
    requirement_docs: List[str]

    features: List[Dict[str, Any]]
    user_stories: List[Dict[str, Any]]
    test_cases: List[Dict[str, Any]]
    selenium_scripts: Dict[str, str]
    validation_results: Dict[str, str]

    collated_docx: str

@telemetry("UI Upload")
def upload_file(uploaded_file):
    if uploaded_file is not None:
        file_path = os.path.join("requirements", uploaded_file.name)
        with open(file_path, "wb") as f:
            f.write(uploaded_file.getbuffer())
        st.success(f"✅ Uploaded {uploaded_file.name} to requirements/")

# -------------------------------
# 2. Helpers to read documents
# -------------------------------
@telemetry("Read Docx")
def read_docx(file_path: str) -> str:
    try:
        doc = DocxDocument(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    except ValueError as e:
        logger.warning(f"Could not read file {file_path}: {e}")
    except FileNotFoundError:
        logger.warning(f"File not found: {file_path}")
    except Exception as e:
        logger.warning(f"An unexpected error occurred while reading {file_path}: {e}")


@telemetry("Read Pdf")
def read_pdf(file_path: str) -> str:
    try:
        text = ""
        with fitz.open(file_path) as pdf:
            for page in pdf:
                text += page.get_text("text") + "\n"
        return text.strip()
    except ValueError as e:
        logger.warning(f"Could not read file {file_path}: {e}")
    except FileNotFoundError:
        logger.warning(f"File not found: {file_path}")
    except Exception as e:
        logger.warning(f"An unexpected error occurred while reading {file_path}: {e}")


@telemetry("Load Requirement Doc")
def load_requirement_docs(folder_path: str) -> List[str]:
    docs = []
    for file_name in os.listdir(folder_path):

        file_path = os.path.join(folder_path, file_name)
        if file_name.lower().endswith(".docx"):
            docs.append(read_docx(file_path))
        elif file_name.lower().endswith(".pdf"):
            docs.append(read_pdf(file_path))
    return docs

    


# -------------------------------
# 3. LLM Setup (Azure OpenAI)
# -------------------------------
# llm = AzureChatOpenAI(
#     deployment_name="your-deployment-name",
#     model_name="gpt-4o",  
#     temperature=0.3,
#     top_p=0.95,
#     openai_api_base=os.getenv("AZURE_OPENAI_ENDPOINT"),
#     openai_api_version="2023-05-15",
#     openai_api_key=os.getenv("AZURE_OPENAI_KEY")
# )


llm = AzureChatOpenAI(
    openai_api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    deployment_name=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
    openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION"),
    temperature=0.23,
    top_p=0.9,
    max_tokens=32000,
)




# -------------------------------
# 4. Nodes
# -------------------------------
@telemetry("Feature Analyzer Node")
def feature_analyzer_node(state: AgentState) -> AgentState:
    docs_text = "\n\n".join(state["requirement_docs"])
    prompt = f"""
    Extract features from the following requirement documents. Reference Output JSON list:
    [{{"id": "F-xxx", "title": "...", "description": "..."}}]

    Documents:
    {docs_text}

    Do not include ```json
    """

    try:
        response = llm.invoke(prompt)   
        features = json.loads(response.content)
        logger.info(f"Extracted {len(features)} features from doc.")
        logger.info(f"Extracted Features => {features}")
        state["features"] = response.content
        return state
    except json.JSONDecodeError as e:
        logger.warning(f"Could not parse features JSON. Skipping. {e}")
    except Exception as e:
        logger.warning(f"Could not parse features JSON. Skipping. {e}")

@telemetry("User Story Node")
def user_story_node(state: AgentState) -> AgentState:

    try:
        all_user_stories = []
        json_feature=json.loads(state['features'])
        for feature in json_feature:
            
            print (f"processing feature = {feature['id']} , {feature}")
            prompt = f"""
            Generate user stories for these features (each feature can have multiple user stories). Reference Output JSON list:
            [{{"id": "F-xxx_US-xxx", "feature_id": "F-xxx", "story": "As a ..."}}]

            Features:
            {feature}

            Do not include ```json
            """
            response = llm.invoke(prompt)
            user_stories=json.loads(response.content)
            all_user_stories.extend(user_stories)
            logger.info(f"Extracted User Stories => {user_stories} user stories from doc.")
            #logger.info(f"Extracted {feature['id']} = {user_stories['id']} user stories from doc.")
        state["user_stories"] = all_user_stories
        logger.info(f"Total Extracted {len(state["user_stories"])} user stories from doc.")
        logger.info(f"Total Extracted User Stories => {state["user_stories"]}")
        return state
    except json.JSONDecodeError as e:
        logger.warning(f"Could not parse user stories JSON. Skipping. {e}")
    except Exception as e:
        logger.warning(f"Could not parse user stories JSON. Skipping. {e.with_traceback}")


@telemetry("Test Case Node")
def test_case_node(state: AgentState) -> AgentState:
    try:
        all_test_cases = []
        #json_user_story=json.loads(state['user_stories'])
        for user_story in state['user_stories']:

            prompt = f"""
            Generate test cases for these user stories (each user stories can have multiple test cases). Reference output JSON list:
            [{{"id": "US-xxx_TC-xxx", "user_story_id": "US-xxx", "steps": ["step1", "step2"], "expected_result": "..."}}]

            User Stories:
            {user_story}

            Generate all the test cases 
            Do not include ```json
            """          
            response = llm.invoke(prompt)
            test_cases=json.loads(response.content)
            all_test_cases.extend(test_cases)
            logger.info(f"Extracted Test Cases => {test_cases}")
        state["test_cases"] = all_test_cases
        logger.info(f"Extracted {len(state["test_cases"])} test cases from doc.")
        logger.info(f"Total Extracted Test Cases => {state["test_cases"]}")
        return state
    except json.JSONDecodeError as e:
        logger.warning(f"Could not parse test cases JSON. Skipping. {e.with_traceback()}")
    except Exception as e:
        logger.warning(f"Could not parse test cases JSON. Skipping. {e.with_traceback()}")


@telemetry("Test Script Node")
def selenium_script_node(state: AgentState) -> AgentState:
    scripts = {}
    try:
        #json_tc = json.loads(state["test_cases"])
        #print(json_tc)
        for tc in state["test_cases"]:
            #print (tc)
            #itc = json.loads(tc)
            prompt = f"""
            Generate a Selenium Python script for this test case and provide full code:

            {tc}

            Filename: GEN-{tc['id']}.py

            Mandatory: 
            - Use best practices and standards
            - Remove comments, notes, summaries, headers, trailers and ```python
            """

            
            response = llm.invoke(prompt)
            filename = f"GEN-{tc['id']}.py"
            filepath = os.path.join(GEN_SCRIPT_FOLDER, filename)
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(response.content)
            scripts[tc["id"]] = filepath

            logger.info(f"Extracted {filename} test scripts from doc.")
        state["selenium_scripts"] = scripts
        logger.info(f"Total test scripts generated: {len(state["selenium_scripts"])}")
        return state
    except Exception as e:
        logger.warning(f"Could not generate test scripts. Skipping. {e.with_traceback()}")



@telemetry("Test Script Validation Node")
def validation_node(state: AgentState) -> AgentState:
    results = {}
    try:
        for tc_id, script_path in state["selenium_scripts"].items():
            with open(script_path, "r", encoding="utf-8") as f:
                script = f.read()
            prompt = f"""
            Validate this Selenium script against best practices and standards. Output "Pass" or "Fail with issues".

            {script}
            """
            response = llm.invoke(prompt)
            results[tc_id] = response.content
            
            state["validation_results"] = results
            logger.info(f"Validation complete for {script_path}")
        return state
    except Exception as e:
            logger.warning(f"Could not validation test scripts. Skipping. {script_path} Error: {e}")

@telemetry("Document Collation Node")
def collation_node(state: AgentState) -> AgentState:
       
    raw_output_path = os.path.join(OUTPUT_FOLDER, GENAI_RAW_OUTPUT)
    output_path = os.path.join(OUTPUT_FOLDER, DOCUMENT_NAME)
    try:

        prompt = f"""
        You are a Business Analyst, please collate Features → User Stories → Test Cases.
        
        Features:
        {state['features']}
        
        Mandatory to include the ids, description, all the steps, expected results.
        """
        
        # prompt = f"""
        # You are a Business Analyst, please collate Features → User Stories → Test Cases.
        
        # Features:
        # {state['features']}
        # User Stories:
        # {state['user_stories']}
        # Test Cases:
        # {state['test_cases']}
        
        # Mandatory to include the ids, description, all the steps, expected results.
        # """
        response = llm.invoke(prompt)

        with open(raw_output_path, "w", encoding="utf-8") as f:
            f.write(response.content)

        markdown_to_word(raw_output_path, output_path)

        print(f"Successfully converted '{raw_output_path}' to '{output_path}'.")

        # doc = DocxDocument()
        # doc.add_heading("Requirement Analysis Report", 0)
        # doc.add_paragraph(response.content)
        
        # doc.save(output_path)
        state["collated_docx"] = output_path
        logger.info(f"Document collation complete for {output_path}")
        return state
    except Exception as e:
        logger.warning(f"Could not collate all documents. Skipping. {e}")


def createLangraphApp():

    # -------------------------------
    # 5. Build Graph
    # -------------------------------
    graph = StateGraph(AgentState)
    graph.add_node("feature_analyzer", feature_analyzer_node)
    graph.add_node("user_story", user_story_node)
    graph.add_node("test_case", test_case_node)
    graph.add_node("selenium_script", selenium_script_node)
    graph.add_node("validation", validation_node)
    graph.add_node("collation", collation_node)

    graph.set_entry_point("feature_analyzer")
    graph.add_edge("feature_analyzer", "user_story")
    graph.add_edge("user_story", "test_case")
    graph.add_edge("test_case", "selenium_script")
    #graph.add_edge("selenium_script", "validation")
    graph.add_edge("selenium_script", "collation")
    graph.add_edge("collation", END)

    # graph.set_entry_point("feature_analyzer")
    # #graph.add_edge("feature_analyzer", "user_story")
    # #graph.add_edge("user_story", "test_case")
    # graph.add_edge("feature_analyzer", "collation")
    # graph.add_edge("collation", END)


    app = graph.compile()

    return app

def small_download_button(file_path, label):
    with open(file_path, "rb") as f:
        data = f.read()
    b64 = base64.b64encode(data).decode()
    href = f'<a href="data:file/txt;base64,{b64}" download="{os.path.basename(file_path)}" class="small-btn">⬇ {label}</a>'
    st.markdown(href, unsafe_allow_html=True)

# Helper function to trigger rerun
def trigger_rerun(key):
    st.session_state[key] = st.session_state.get(key, 0) + 1

def st_initialize():
    st.set_page_config(
            page_title="AI Test Script Generator",
            page_icon="🤖",
            layout="centered"
     )
    
    st.title("🤖 AI Test Script Generator")
    st.write("Upload your requirements document and generate Selenium test scripts automatically.")

def st_upload_file() -> List:
    uploaded_file = st.file_uploader("Upload Requirements Document", type=["pdf", "docx"])
    upload_file(uploaded_file)
    docs = load_requirement_docs(INPUT_REQ_DOCS_FOLDER)
    return docs

def st_sidebar():

    # Custom CSS for smaller expanders + buttons
    st.markdown(
        """
        <style>
        /* Shrink expander title */
        div.streamlit-expanderHeader {
            font-size: 0.2rem !important;
            padding: 2px 4px !important;
        }

        /* Shrink expander content */
        div.streamlit-expanderContent {
            padding: 4px 6px !important;
        }

        /* Compact custom download button */
        .small-btn {
            display: inline-block;
            padding: 3px 8px;
            margin: 2px 0;
            background-color: none;
            color: white;
            text-decoration: none;
            font-size: 0.75rem;
            border-radius: 4px;
        }
        .small-btn:hover {
            background-color: #f5f4ed;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Ensure session_state keys exist
    for k in ["refresh_documents", "refresh_scripts"]:
        if k not in st.session_state:
            st.session_state[k] = 0
    with st.sidebar:

        if os.path.exists(OUTPUT_FOLDER):
            with st.expander("✅ Generated Documents", expanded=False): 
                if st.button("🔄 Refresh", key="refresh_documents_btn"):
                    trigger_rerun("refresh_documents")
                col1, col2 = st.columns([2, 1])
                with col1:
                    delete_click = st.button("🗑️ Delete All Documents")
                with col2:
                    delete_confirm = st.checkbox("Confirm", key="doc_delete_confirm_checkbox")  # inline checkbox
                #delete_confirm = st.checkbox("Confirm delete all outputs")
                if delete_click and delete_confirm:
                    try:
                        for fname in os.listdir(OUTPUT_FOLDER):
                            fpath = os.path.join(OUTPUT_FOLDER, fname)
                            os.remove(fpath)
                        st.success("✅ All generated files have been deleted.")
                    except Exception as e:
                        st.error(f"❌ Error deleting files: {e}")

                for fname in os.listdir(OUTPUT_FOLDER):
                    fpath = os.path.join(OUTPUT_FOLDER, fname) 
                    small_download_button(fpath, fname)

        if os.path.exists(GEN_SCRIPT_FOLDER):
            with st.expander("✅ Generated Selenium Scripts", expanded=False): 
                if st.button("🔄 Refresh", key="refresh_scripts_btn"):
                    trigger_rerun("refresh_scripts")
                col1, col2 = st.columns([2, 1])
                with col1:
                    delete_click = st.button("🗑️ Delete All Scripts")
                with col2:
                    delete_confirm = st.checkbox("Confirm", key="scripts_delete_confirm_checkbox")  # inline checkbox
                if delete_click and delete_confirm:
                    try:
                        for fname in os.listdir(GEN_SCRIPT_FOLDER):
                            fpath = os.path.join(GEN_SCRIPT_FOLDER, fname)
                            os.remove(fpath)
                        st.success("✅ All generated scripts have been deleted.")
                    except Exception as e:
                        st.error(f"❌ Error deleting scripts: {e}")

                for fname in os.listdir(GEN_SCRIPT_FOLDER):
                    fpath = os.path.join(GEN_SCRIPT_FOLDER, fname) 
                    small_download_button(fpath, fname)



def st_start_processing(app) -> AgentState:

    if st.button("🚀 Generate Test Scripts"):
        st.info("⚡ Running pipeline... Logs will appear below.")
        
        # Create a log placeholder in UI
        log_placeholder = st.empty()
        streamlit_handler = StreamlitLogger(log_placeholder)
        formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")
        streamlit_handler.setFormatter(formatter)

        # Attach handler to root logger
        stlogger = logging.getLogger()
        stlogger.setLevel(logging.INFO)
        stlogger.addHandler(streamlit_handler)


        with st.spinner("Processing..."):
            try:
                init_state: AgentState = {"requirement_docs": docs}
                final_state = app.invoke(init_state)
                st.success("🎉 Pipeline finished successfully!")
                st.balloons()
                return final_state
            except Exception as e:
                st.error(f"❌ Error: {str(e)}") 

# -------------------------------
# 6. Run Example 

# -------------------------------
if __name__ == "__main__":
  
    INPUT_REQ_DOCS_FOLDER = "requirement_docs"
    GEN_SCRIPT_FOLDER = "GEN-TESTSCRIPTS"
    OUTPUT_FOLDER = "generated_outputs"
    GENAI_RAW_OUTPUT="GenAI_RAW_OUTPUT.md"
    DOCUMENT_NAME="GenAI_Features_UserStories_Tescases.docx"
    os.makedirs(GEN_SCRIPT_FOLDER, exist_ok=True)
    os.makedirs(OUTPUT_FOLDER, exist_ok=True)
    os.makedirs(INPUT_REQ_DOCS_FOLDER, exist_ok=True)
    st_initialize()
    docs=st_upload_file()
    app=createLangraphApp()
    final_state=st_start_processing(app)
    st_sidebar() 
    #logger.info("Agent Execution complete!")

